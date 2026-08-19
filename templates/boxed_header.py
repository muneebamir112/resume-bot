# -*- coding: utf-8 -*-
"""Template: Boxed Header — centered header block, section titles inside
bordered rectangle boxes, plain black-on-white body text throughout."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .common import (
    NAME, EMAIL, PHONE, ADDRESS, EXPERIENCE, BLACK, GRAY,
    set_margins, set_base_font, set_cell_margins, set_table_box_borders,
)

PAGE_WIDTH = Inches(7.2)


def _boxed_header(doc, text):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(10)
    run = spacer.add_run("")
    run.font.size = Pt(4)

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_box_borders(table, color="000000", sz=8)
    cell = table.cell(0, 0)
    cell.width = PAGE_WIDTH
    set_cell_margins(cell, left=140, right=140, top=70, bottom=70)
    p = cell.paragraphs[0]
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = BLACK
    return table


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def _role_block(doc, title, company, location, dates):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(0)
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = BLACK

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(4)
    run2 = p_meta.add_run(f"{company} | {location} | {dates}")
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = GRAY


def build(data: dict, output_path: str):
    doc = Document()
    set_base_font(doc)
    set_margins(doc)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(4)
    run = name_p.add_run(NAME)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BLACK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    run = subtitle_p.add_run(data["subtitle"])
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(4)
    run = contact_p.add_run(
        f"{EMAIL} | {PHONE} | {ADDRESS} | {EXPERIENCE} Experience"
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK

    _boxed_header(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(6)
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.line_spacing = 1.12
    run = sp.add_run(data["summary"])
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    _boxed_header(doc, "Work Experience")
    for role in data["experience"]:
        _role_block(doc, role["title"], role["company"], role["location"], role["dates"])
        for b in role["bullets"]:
            _bullet(doc, b)

    _boxed_header(doc, "Skills")
    for label, value in data["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6) if label == data["skills"][0][0] else Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run2 = p.add_run(value)
        run2.font.size = Pt(11)
        run2.font.color.rgb = BLACK

    _boxed_header(doc, "Education")
    edu_p1 = doc.add_paragraph()
    edu_p1.paragraph_format.space_before = Pt(6)
    edu_p1.paragraph_format.space_after = Pt(0)
    run = edu_p1.add_run(data["education"]["school"])
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    edu_p2 = doc.add_paragraph()
    edu_p2.paragraph_format.space_after = Pt(4)
    run = edu_p2.add_run(data["education"]["degree_line"])
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    _boxed_header(doc, "Keywords")
    kp = doc.add_paragraph()
    kp.paragraph_format.space_before = Pt(6)
    kp.paragraph_format.line_spacing = 1.12
    run = kp.add_run(data["keywords"])
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    doc.save(output_path)
    return output_path
