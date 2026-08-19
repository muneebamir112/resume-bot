# -*- coding: utf-8 -*-
"""Template: Underline Header — centered header with stacked, labeled
contact lines; bold black section headers with a thin underline;
single-column skills list. Black & white throughout."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .common import (
    NAME, EMAIL, PHONE, ADDRESS, EXPERIENCE, BLACK, GRAY,
    add_bottom_border, set_margins, set_base_font,
)


def _section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = BLACK
    add_bottom_border(p, color="000000", sz=6)
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK


def _role_block(doc, title, company, location, dates):
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(0)
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(4)
    run2 = p_meta.add_run(f"{company} | {location} | {dates}")
    run2.bold = True
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = BLACK


def _contact_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    return p


def build(data: dict, output_path: str):
    doc = Document()
    set_base_font(doc)
    set_margins(doc)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    run = name_p.add_run(NAME)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = BLACK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(6)
    run = subtitle_p.add_run(data["subtitle"].upper())
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY

    _contact_line(doc, f"Email: {EMAIL}")
    _contact_line(doc, f"Phone: {PHONE}")
    _contact_line(doc, f"Address: {ADDRESS}")
    last = _contact_line(doc, f"Experience: {EXPERIENCE}")
    last.paragraph_format.space_after = Pt(2)
    add_bottom_border(last, color="000000", sz=10)

    _section_header(doc, "Professional Summary")
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    sp.paragraph_format.line_spacing = 1.12
    run = sp.add_run(data["summary"])
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    _section_header(doc, "Work Experience")
    for role in data["experience"]:
        _role_block(doc, role["title"], role["company"], role["location"], role["dates"])
        for b in role["bullets"]:
            _bullet(doc, b)

    _section_header(doc, "Skills")
    for label, value in data["skills"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run2 = p.add_run(value)
        run2.font.size = Pt(11)
        run2.font.color.rgb = BLACK

    _section_header(doc, "Education")
    edu_p1 = doc.add_paragraph()
    edu_p1.paragraph_format.space_after = Pt(0)
    run = edu_p1.add_run(data["education"]["school"])
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    edu_p2 = doc.add_paragraph()
    edu_p2.paragraph_format.space_after = Pt(4)
    run = edu_p2.add_run(data["education"]["degree_line"])
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    _section_header(doc, "Keywords")
    kp = doc.add_paragraph()
    kp.paragraph_format.line_spacing = 1.12
    run = kp.add_run(data["keywords"])
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    doc.save(output_path)
    return output_path
